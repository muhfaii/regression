"""Export endpoints: /api/export/* and public /api/share/*"""
from __future__ import annotations

import dataclasses
import io
import uuid
from datetime import datetime

from docx import Document
from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse

from backend.services import session_store

router = APIRouter(tags=["export"])

_WORD_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

_SKIP_STAT_KEYS = frozenset({
    "variables", "groups", "coefficients", "post_hoc",
    "contingency_table", "outcome_categories",
})


# ---------------------------------------------------------------------------
# Endpoints
# ---------------------------------------------------------------------------

@router.post("/api/export/word/{result_id}")
async def export_word(result_id: str, session_id: str) -> StreamingResponse:
    raw = session_store.get_result(session_id, result_id)
    if raw is None:
        raise HTTPException(status_code=404, detail="Result not found.")

    result = _normalize(raw)
    doc = _build_docx(result)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"{result.get('test_key', 'result')}_{result_id[:8]}.docx"
    return StreamingResponse(
        buf,
        media_type=_WORD_MEDIA_TYPE,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/api/export/share/{result_id}")
async def create_share_link(result_id: str, session_id: str) -> dict:
    raw = session_store.get_result(session_id, result_id)
    if raw is None:
        raise HTTPException(status_code=404, detail="Result not found.")

    token = str(uuid.uuid4())
    await session_store.save_share(token, raw)
    return {"token": token, "url": f"/share/{token}"}


@router.get("/api/share/{token}")
async def get_shared_result(token: str) -> dict:
    raw = await session_store.get_share(token)
    if raw is None:
        raise HTTPException(status_code=404, detail="Share link not found.")
    return _normalize(raw)


# ---------------------------------------------------------------------------
# Reproducibility log
# ---------------------------------------------------------------------------

@router.get("/api/export/log/{session_id}")
async def get_session_log(session_id: str) -> dict:
    session = session_store.get_session(session_id)
    if session is None:
        raise HTTPException(status_code=404, detail="Session not found or expired.")
    return {"filename": session.filename, "steps": session.steps}


@router.get("/api/export/log/{session_id}/download")
async def download_session_log(session_id: str) -> StreamingResponse:
    session = session_store.get_session(session_id)
    if session is None:
        raise HTTPException(status_code=404, detail="Session not found or expired.")

    text = _build_log_text(session)
    buf = io.BytesIO(text.encode("utf-8"))
    return StreamingResponse(
        buf,
        media_type="text/plain",
        headers={"Content-Disposition": f'attachment; filename="analysis_log_{session_id[:8]}.txt"'},
    )


# ---------------------------------------------------------------------------
# Multi-result report
# ---------------------------------------------------------------------------

@router.post("/api/export/report/{session_id}")
async def export_report(session_id: str, body: dict | None = None) -> StreamingResponse:
    session = session_store.get_session(session_id)
    if session is None:
        raise HTTPException(status_code=404, detail="Session not found or expired.")

    result_ids = (body or {}).get("result_ids") or list(session.results.keys())
    results = [_normalize(session.results[rid]) for rid in result_ids if rid in session.results]
    if not results:
        raise HTTPException(status_code=422, detail="No results to include in the report.")

    doc = _build_report_docx(session.filename, results)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"analysis_report_{session_id[:8]}.docx"
    return StreamingResponse(
        buf,
        media_type=_WORD_MEDIA_TYPE,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ---------------------------------------------------------------------------
# Session-level (multi-result) sharing
# ---------------------------------------------------------------------------

@router.post("/api/export/share-session/{session_id}")
async def create_session_share_link(session_id: str, body: dict | None = None) -> dict:
    session = session_store.get_session(session_id)
    if session is None:
        raise HTTPException(status_code=404, detail="Session not found or expired.")

    result_ids = (body or {}).get("result_ids") or list(session.results.keys())
    results = [_normalize(session.results[rid]) for rid in result_ids if rid in session.results]
    if not results:
        raise HTTPException(status_code=422, detail="No results to share.")

    token = str(uuid.uuid4())
    session_store.save_share_session(token, {"filename": session.filename, "results": results})
    return {"token": token, "url": f"/share/session/{token}"}


@router.get("/api/share/session/{token}")
async def get_shared_session(token: str) -> dict:
    bundle = session_store.get_share_session(token)
    if bundle is None:
        raise HTTPException(status_code=404, detail="Share link not found.")
    return bundle


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _normalize(result: dict) -> dict:
    """Convert nested dataclass instances in a stored result dict to plain dicts."""
    out: dict = {}
    for k, v in result.items():
        if dataclasses.is_dataclass(v) and not isinstance(v, type):
            out[k] = dataclasses.asdict(v)
        elif isinstance(v, list):
            out[k] = [
                dataclasses.asdict(i) if dataclasses.is_dataclass(i) and not isinstance(i, type) else i
                for i in v
            ]
        else:
            out[k] = v
    return out


def _fmt(val: object) -> str:
    if isinstance(val, float):
        return f"{val:.4f}"
    if isinstance(val, bool):
        return "Yes" if val else "No"
    return str(val)


def _add_result_section(doc: Document, result: dict, title_level: int = 0) -> None:
    """Append one result's content to doc, starting headings at title_level."""
    doc.add_heading(result.get("test_name", "Analysis"), title_level)
    doc.add_paragraph(f"N = {result.get('n_obs', '—')}")

    # Key scalar statistics table
    stats = result.get("statistics", {})
    scalar_stats = {
        k: v for k, v in stats.items()
        if k not in _SKIP_STAT_KEYS and isinstance(v, (int, float, str, bool))
    }
    if scalar_stats:
        doc.add_heading("Statistics", level=title_level + 1)
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "Statistic"
        hdr[1].text = "Value"
        for key, val in scalar_stats.items():
            row = tbl.add_row().cells
            row[0].text = key.replace("_", " ").title()
            row[1].text = _fmt(val)

    # Effect size
    effect = result.get("effect_size")
    if effect:
        doc.add_heading("Effect Size", level=title_level + 1)
        doc.add_paragraph(
            f"{effect['name']}: {effect['value']:.3f} ({effect['interpretation']})"
        )

    # Assumption checks
    checks = result.get("assumption_checks", [])
    if checks:
        doc.add_heading("Assumption Checks", level=title_level + 1)
        icons = {"pass": "✓", "amber": "⚠", "fail": "✗"}
        for check in checks:
            icon = icons.get(check.get("status", ""), "")
            doc.add_paragraph(f"{icon} {check['name']}: {check['detail']}")
            if check.get("fix_suggestion"):
                doc.add_paragraph(f"Suggestion: {check['fix_suggestion']}")

    # Interpretation — APA first (most useful for export), then plain English
    interp = result.get("interpretation", {})
    if interp:
        doc.add_heading("Interpretation", level=title_level + 1)
        doc.add_heading("APA 7", level=title_level + 2)
        doc.add_paragraph(interp.get("apa", ""))
        doc.add_heading("Plain English", level=title_level + 2)
        doc.add_paragraph(interp.get("plain", ""))


def _build_docx(result: dict) -> Document:
    doc = Document()
    _add_result_section(doc, result, title_level=0)
    return doc


def _build_report_docx(filename: str, results: list[dict]) -> Document:
    doc = Document()
    doc.add_heading("Analysis Report", 0)
    doc.add_paragraph(f"Dataset: {filename}")
    doc.add_paragraph(f"{len(results)} result(s) included")
    for i, result in enumerate(results):
        if i > 0:
            doc.add_page_break()
        _add_result_section(doc, result, title_level=1)
    return doc


def _build_log_text(session) -> str:
    lines = [f"Analysis log — {session.filename}", ""]
    for step in session.steps:
        ts = datetime.fromtimestamp(step["timestamp"]).strftime("%Y-%m-%d %H:%M:%S")
        lines.append(f"[{ts}] {step['action']}: {step['detail']}")
    return "\n".join(lines)
